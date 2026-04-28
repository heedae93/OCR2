"""
NAR-DB구축-008 요구사항 대비 현황 보고서 생성
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), kwargs.get(edge, 'single'))
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:color'), 'CCCCCC')
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_cell_text(cell, text, bold=False, size=9, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[i])

doc = Document()

# 기본 여백 설정
section = doc.sections[0]
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ───────────────────────────────────────────────
# 제목
# ───────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('NAR-DB구축-008 요구사항 대비\nMASKING_OCR 구현 현황 분석')
run.font.size = Pt(18)
run.font.bold = True
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_p.add_run(f'작성일: {datetime.date.today().strftime("%Y년 %m월 %d일")}')
date_run.font.size = Pt(10)
date_run.font.name = '맑은 고딕'
date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()

# ───────────────────────────────────────────────
# 1. 개요
# ───────────────────────────────────────────────
def add_section_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)

def add_sub_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.color.rgb = RGBColor(0x37, 0x37, 0x37)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)

def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p.paragraph_format.space_after = Pt(4)

add_section_title(doc, '1. 개요')
add_body(doc,
    '본 문서는 국회기록물 DB구축 요구사항 NAR-DB구축-008 (XML·JSON 파일 생성, AI OCR 기술 적용, LLM·VLM 활용)에 대해 '
    '현재 MASKING_OCR(BBOCR) 프로젝트의 구현 현황을 분석하고, '
    '보강이 필요한 항목과 우선순위를 정리한 보고서입니다.'
)

# 요구사항 요약 표
add_sub_title(doc, '■ 분석 대상 요구사항')
req_table = doc.add_table(rows=5, cols=2)
req_table.style = 'Table Grid'
headers = [('요구사항 고유번호', 'NAR-DB구축-008'),
           ('요구사항 분류', '국회기록물 DB구축 요구사항'),
           ('요구사항 명칭', 'XML·JSON 파일 생성'),
           ('주요 기술 항목', 'AI OCR 기술 적용 / LLM·VLM 활용 / XML·JSON 생성'),
           ('납품 품질 기준', 'XSD 유효성 검증 100% / 필수 메타데이터 입력률 100%')]
for i, (k, v) in enumerate(headers):
    set_cell_bg(req_table.rows[i].cells[0], 'EFF6FF')
    add_cell_text(req_table.rows[i].cells[0], k, bold=True, size=9)
    add_cell_text(req_table.rows[i].cells[1], v, size=9)
set_col_widths(req_table, [5.0, 11.0])

doc.add_paragraph()

# ───────────────────────────────────────────────
# 2. 전체 현황 요약
# ───────────────────────────────────────────────
add_section_title(doc, '2. 전체 구현 현황 요약')

summary_table = doc.add_table(rows=4, cols=3)
summary_table.style = 'Table Grid'
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers_row = summary_table.rows[0]
for cell, text in zip(headers_row.cells, ['구분', '항목 수', '세부 내용']):
    set_cell_bg(cell, '1A56DB')
    add_cell_text(cell, text, bold=True, size=10, color=(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER)

rows_data = [
    ('✅ 구현 완료', '13개', 'AI OCR, 레이아웃 분석, 표 인식, 다단 인식, 읽기순서,\nPII 마스킹, JSON/XML 생성, NER, Searchable PDF, 메타데이터 관리 등'),
    ('⚠️ 부분 구현', '3개', '직인·서명 인식(비활성화), LLM 기반 추출(규칙 기반 대체),\nXML 정합성 검증'),
    ('❌ 미구현', '5개', '국제표준 메타데이터(Dublin Core/Schema.org),\nXSD 스키마·유효성 검증, 문서 요약, 주제어 태깅, 목차 추출'),
]
colors = ['E8F5E9', 'FFF8E1', 'FFEBEE']
for i, (label, count, detail) in enumerate(rows_data, 1):
    set_cell_bg(summary_table.rows[i].cells[0], colors[i-1])
    add_cell_text(summary_table.rows[i].cells[0], label, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(summary_table.rows[i].cells[1], count, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(summary_table.rows[i].cells[2], detail, size=9)

set_col_widths(summary_table, [3.5, 2.0, 10.5])
doc.add_paragraph()

# ───────────────────────────────────────────────
# 3. 구현 완료 항목 상세
# ───────────────────────────────────────────────
add_section_title(doc, '3. 구현 완료 항목 상세')

done_items = [
    ('AI OCR 텍스트 인식',
     'PaddleOCR + PPStructure V3 기반 다중 엔진 구성, 타일링 처리로 대용량 이미지 지원',
     'core/ocr_engine.py\ncore/pp_structure_engine.py'),
    ('한국어 특화 OCR',
     '중국어 계열 감지 모델(ch_PP-OCRv4_det) + 커스텀 한국어 인식 모델 조합, GPU/CPU 모드 지원',
     'core/ocr_engine.py'),
    ('문서 레이아웃 분석',
     'PP-DocLayout-L 기반 제목/본문/표/이미지/수식/참조 영역 자동 분리, OCR 블록과 레이아웃 매핑',
     'core/layout_detector.py'),
    ('표(테이블) 인식',
     'SLANet+ 기반 표 구조 추출, Microsoft Table Transformer 엔진 선택 지원, HTML 변환 가능',
     'core/pp_structure_engine.py\ncore/table_transformer_engine.py'),
    ('다단(멀티컬럼) 인식',
     '간격 분석 기반 2단 레이아웃 자동 감지, 신뢰도 점수 및 경계 좌표 산출',
     'core/column_detector.py'),
    ('읽기순서 정렬',
     '행 클러스터링 기반 좌→우, 위→아래 정렬, 2단 레이아웃 컬럼별 정렬 지원',
     'core/reading_order_sorter.py'),
    ('개인정보(PII) 감지·마스킹',
     '이름/전화/주민번호/주소/이메일/계좌 등 14종 감지, KoBERT NER 연동, 마스킹 PDF 자동 생성',
     'core/pii_extractor.py\napi/masking.py'),
    ('JSON 파일 생성',
     'OCR 결과·레이아웃·마스킹 정보를 UTF-8 JSON으로 job 단위 저장',
     'api/ocr.py → {job_id}_ocr.json'),
    ('XML 파일 생성',
     'ABBYY FineReader 호환 XML 생성, UTF-8 인코딩, 페이지·블록 구조 포함',
     'api/export.py → create_abbyy_xml()'),
    ('NER 개체명 인식',
     'KoELECTRA/KoBERT 기반 인물·기관·장소·날짜·시간 개체명 자동 태깅',
     'utils/ner_extractor.py\ncore/pii_extractor.py'),
    ('Searchable PDF 생성',
     '보이지 않는 텍스트 레이어 삽입 방식, 한글 폰트 지원, 정밀 좌표 스케일링',
     'core/invisible_layer.py\ncore/text_scaler.py'),
    ('메타데이터 관리',
     '문서명·날짜·키워드·문서유형·언어 추출, DB 저장, 메타데이터 V3 API 제공',
     'api/metadata_v3.py\nutils/metadata_extractor.py'),
    ('세션·작업 단위 관리',
     '기록물 건 단위 세션/작업 관리, DB 기반 상태 추적, Celery 비동기 처리',
     'api/sessions.py\nutils/job_manager.py'),
]

done_table = doc.add_table(rows=len(done_items)+1, cols=3)
done_table.style = 'Table Grid'

for cell, text in zip(done_table.rows[0].cells, ['기능', '구현 내용', '관련 파일']):
    set_cell_bg(cell, '1E4D2B')
    add_cell_text(cell, text, bold=True, size=9, color=(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER)

for i, (func, detail, files) in enumerate(done_items, 1):
    bg = 'F0FDF4' if i % 2 == 0 else 'FFFFFF'
    set_cell_bg(done_table.rows[i].cells[0], bg)
    set_cell_bg(done_table.rows[i].cells[1], bg)
    set_cell_bg(done_table.rows[i].cells[2], bg)
    add_cell_text(done_table.rows[i].cells[0], func, bold=True, size=8)
    add_cell_text(done_table.rows[i].cells[1], detail, size=8)
    add_cell_text(done_table.rows[i].cells[2], files, size=8)

set_col_widths(done_table, [3.5, 8.0, 4.5])
doc.add_paragraph()

# ───────────────────────────────────────────────
# 4. 부분 구현 항목
# ───────────────────────────────────────────────
add_section_title(doc, '4. 부분 구현 항목 — 보강 필요')

partial_items = [
    ('직인·서명 영역 인식',
     '`use_seal_recognition=False`로 비활성화 상태.\n코드 내 옵션 존재하나 실제 동작 안 함.',
     'core/pp_structure_engine.py\n(use_seal_recognition 파라미터)',
     '① use_seal_recognition=True 활성화\n② 인식 결과를 OCR 결과·JSON·XML에 반영\n③ 직인 영역 별도 레이아웃 타입으로 표시'),
    ('LLM·VLM 기반 정보 추출',
     'KoBERT NER 수준의 개체명 인식만 존재.\n발신/수신기관, 문서번호 자동추출 미구현.\nLLM API 연동 없음.',
     'core/pii_extractor.py\nutils/metadata_extractor.py',
     '① 발신/수신기관 규칙 기반 추출기 추가\n② 문서번호 패턴 추출기 추가\n③ Claude/GPT API 연동으로 요약·태깅 지원'),
    ('XML 구조 정합성 검증',
     '생성된 XML이 PDF 페이지 구조·목차와\n일치하는지 자동 검증 로직 없음.',
     'api/export.py',
     '① XML ↔ PDF 페이지 수 일치 검증\n② XML ↔ JSON 주요 메타데이터 비교 로직 추가\n③ 정합성 검증 리포트 생성'),
]

partial_table = doc.add_table(rows=len(partial_items)+1, cols=4)
partial_table.style = 'Table Grid'

for cell, text in zip(partial_table.rows[0].cells, ['기능', '현재 상태', '관련 파일', '보강 방안']):
    set_cell_bg(cell, 'B45309')
    add_cell_text(cell, text, bold=True, size=9, color=(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER)

for i, (func, status, files, plan) in enumerate(partial_items, 1):
    bg = 'FFFBEB' if i % 2 == 0 else 'FEF3C7'
    for cell in partial_table.rows[i].cells:
        set_cell_bg(cell, bg)
    add_cell_text(partial_table.rows[i].cells[0], func, bold=True, size=8)
    add_cell_text(partial_table.rows[i].cells[1], status, size=8)
    add_cell_text(partial_table.rows[i].cells[2], files, size=8)
    add_cell_text(partial_table.rows[i].cells[3], plan, size=8)

set_col_widths(partial_table, [2.8, 4.2, 3.5, 5.5])
doc.add_paragraph()

# ───────────────────────────────────────────────
# 5. 미구현 항목
# ───────────────────────────────────────────────
add_section_title(doc, '5. 미구현 항목 — 신규 개발 필요')

new_items = [
    ('국제표준 메타데이터\n(Dublin Core / Schema.org)',
     '납품 XML·JSON에 국제표준 필드 구조 없음.\n착수보고서에서 확정될 필수 항목 미반영.',
     '높음',
     '① Dublin Core 15개 핵심 요소 매핑 정의\n② Schema.org 스키마 적용\n③ XML·JSON 출력에 표준 필드 포함'),
    ('XSD 스키마 파일 및 유효성 검증',
     'XML Schema(XSD) 파일 자체 없음.\n생성 XML의 XSD 검증 통과율 100% 요건 미충족.',
     '높음',
     '① 납품 XSD 스키마 파일 작성\n② 생성된 XML을 lxml로 XSD 검증\n③ 검증 실패 시 오류 리포트 생성'),
    ('XML·JSON 정합성 자동 검증',
     'XML과 JSON 간 구조·메타데이터 일치 여부\n자동 확인 기능 없음.',
     '높음',
     '① XML ↔ JSON 주요 필드 비교 함수 작성\n② 불일치 항목 리포트 자동 생성\n③ 납품 전 최종 검증 파이프라인 구성'),
    ('문서 핵심 내용 요약 생성',
     'LLM 기반 자동 요약 기능 없음.',
     '중간',
     '① Claude API / 오픈소스 LLM 연동\n② 문서 요약 결과를 메타데이터에 저장\n③ JSON 출력에 summary 필드 추가'),
    ('주제어·개체명 자동 태깅',
     'TF-IDF 키워드는 있으나 LLM 기반\n주제어 태깅 및 사건 개체명 없음.',
     '중간',
     '① LLM 기반 주제어 추출 연동\n② 인물·기관·사건 태깅 결과를 JSON에 저장\n③ 검색 서비스 연동 필드로 구조화'),
]

new_table = doc.add_table(rows=len(new_items)+1, cols=4)
new_table.style = 'Table Grid'

for cell, text in zip(new_table.rows[0].cells, ['기능', '현재 상태', '우선순위', '개발 방안']):
    set_cell_bg(cell, '7F1D1D')
    add_cell_text(cell, text, bold=True, size=9, color=(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER)

priority_colors = {'높음': 'FEE2E2', '중간': 'FEF9C3'}
for i, (func, status, priority, plan) in enumerate(new_items, 1):
    bg = priority_colors.get(priority, 'FFFFFF')
    for cell in new_table.rows[i].cells:
        set_cell_bg(cell, bg)
    add_cell_text(new_table.rows[i].cells[0], func, bold=True, size=8)
    add_cell_text(new_table.rows[i].cells[1], status, size=8)
    p_color = (0xB9, 0x1C, 0x1C) if priority == '높음' else (0x85, 0x70, 0x0A)
    add_cell_text(new_table.rows[i].cells[2], priority, bold=True, size=9, color=p_color, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(new_table.rows[i].cells[3], plan, size=8)

set_col_widths(new_table, [3.0, 4.0, 1.8, 7.2])
doc.add_paragraph()

# ───────────────────────────────────────────────
# 6. 보강 우선순위 로드맵
# ───────────────────────────────────────────────
add_section_title(doc, '6. 보강 우선순위 로드맵')

road_table = doc.add_table(rows=6, cols=4)
road_table.style = 'Table Grid'

for cell, text in zip(road_table.rows[0].cells, ['순위', '항목', '분류', '기대 효과']):
    set_cell_bg(cell, '1E3A5F')
    add_cell_text(cell, text, bold=True, size=9, color=(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER)

road_data = [
    ('1', 'XSD 스키마 정의 + XML 유효성 검증', '미구현', '납품 품질 기준 100% 충족'),
    ('2', 'Dublin Core / Schema.org 표준 메타데이터', '미구현', '국제표준 준수 · 납품 요건 충족'),
    ('3', 'XML·JSON 정합성 자동 검증 파이프라인', '미구현', '납품 전 품질 자동 보증'),
    ('4', '발신/수신기관·문서번호 자동추출', '부분구현', '메타데이터 완성도 향상'),
    ('5', '직인·서명 영역 인식 활성화', '부분구현', '문서 구조 인식 정확도 향상'),
]

row_bg = ['EFF6FF', 'F0FDF4', 'FFF8E1', 'FEF3C7', 'FFF1F2']
for i, (rank, item, cls, effect) in enumerate(road_data, 1):
    bg = row_bg[i-1]
    for cell in road_table.rows[i].cells:
        set_cell_bg(cell, bg)
    add_cell_text(road_table.rows[i].cells[0], rank, bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(road_table.rows[i].cells[1], item, bold=True, size=9)
    add_cell_text(road_table.rows[i].cells[2], cls, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(road_table.rows[i].cells[3], effect, size=9)

set_col_widths(road_table, [1.2, 6.0, 2.5, 6.3])
doc.add_paragraph()

# ───────────────────────────────────────────────
# 7. 결론
# ───────────────────────────────────────────────
add_section_title(doc, '7. 결론')
add_body(doc,
    '현재 MASKING_OCR(BBOCR) 프로젝트는 AI OCR 파이프라인의 핵심 기능(텍스트 인식, 레이아웃 분석, '
    '표 인식, 다단 처리, PII 마스킹, Searchable PDF 생성)이 잘 구현되어 있어 '
    'NAR-DB구축-008 요구사항의 기술 기반은 충분히 갖추고 있습니다.'
)
add_body(doc,
    '다만 납품 품질 기준(XSD 유효성 검증 100%, 국제표준 메타데이터)과 '
    'LLM 기반 고급 추출(발신/수신기관, 요약, 주제어 태깅) 부분이 보강 대상이며, '
    '이 항목들을 중심으로 개발 일정을 수립할 것을 권장합니다.'
)

# 저장
output_path = r'd:\project\MASKING_OCR\NAR-DB구축-008_구현현황분석.docx'
doc.save(output_path)
print(f'문서 생성 완료: {output_path}')
